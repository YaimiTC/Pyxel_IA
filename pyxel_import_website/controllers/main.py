# Part of Odoo. See LICENSE file for full copyright and licensing details

import html as html_lib
import json
import logging
import base64
import re
from datetime import datetime

from dateutil.relativedelta import relativedelta
from odoo import http, _
from odoo.addons.website.controllers import form
from odoo.addons.website_crm.controllers.website_form import WebsiteForm as WebsiteForm2
from odoo.exceptions import ValidationError
from odoo.http import Stream, request, Response
from odoo.addons.base.models.ir_qweb_fields import nl2br_enclose

_logger = logging.getLogger(__name__)

DOC_FIELD_LABELS = {
    'doc_mipyme_escritura_notarial': 'Escritura notarial',
    'doc_mipyme_registro_mercantil': 'Registro mercantil',
    'doc_mipyme_carnet_nit': 'Carnet con el NIT',
    'doc_mipyme_contrato_banco': 'Contrato de banco CUP, MLC o USD o certifico bancario',
    'doc_mipyme_certifico_adeudo': 'Certifico de no adeudo o comprobante del último pago',
    'doc_sucursal_escrituras_constitucion': 'Escrituras de constituciones y modificaciones efectuadas',
    'doc_sucursal_registro_mercantil': 'Inscripción en el registro mercantil',
    'doc_sucursal_licencia_camara': 'Licencia de Cámara de Comercio',
    'doc_sucursal_planilla_contribuyente': 'Planilla de inscripción o actualización en el registro de contribuyente',
    'doc_sucursal_contrato_banco': 'Contrato de banco CUP, MLC o USD o certifico bancario',
    'doc_sucursal_resolucion_mincex': 'Resolución del Ministerio del Comercio Exterior y la Inversión Extranjera',
    'doc_estatal_resoluciones': 'Resoluciones constitutivas',
    'doc_estatal_reup': 'Documento acreditativo del Código REUP',
    'doc_estatal_nit': 'Documento acreditativo del NIT',
    'doc_estatal_contrato_banco': 'Contrato de banco CUP, MLC o USD o certifico bancario',
    'doc_cna_onat': 'Documento acreditativo ONAT',
    'doc_cna_reane': 'Documento acreditativo Código REANE',
    'doc_cna_escritura_notarial': 'Escritura notarial',
    'doc_cna_registro_mercantil': 'Registro mercantil',
    'doc_cna_carnet_nit': 'Carnet con el NIT',
    'doc_cna_contrato_banco': 'Contrato de banco CUP, MLC o USD o certifico bancario',
    'doc_cna_certifico_adeudo': 'Certifico de no adeudo o comprobante del último pago',
    'doc_prov_escritura': 'Escritura de constitución y modificaciones efectuadas',
    'doc_prov_registro_mercantil': 'Inscripción en el registro mercantil',
    'doc_prov_poder_acreditativo': 'Poder acreditativo de legitimación de representantes',
    'doc_prov_certifico_bancario': 'Certifico de cuenta bancaria con el que va a operar el contrato',
    'doc_prov_codigo_mincex': 'Código MIncex',
    'doc_adicional': 'Documentación adicional',
}

operation_banner_img = {
    "logistic": "banner_logistica_3.svg",
    "import": "banner_importaciones_2.svg",
    "accreditation": "acreditacion_2.svg",
    "accreditation_business": "acreditacion_2.svg",
}
SPG = 20  # Proveedores Per Page
SPR = 4  # Proveedores Per Row


def get_render_values(kw):
    country = request.env["res.country"].sudo()
    contact_types = request.env["res.partner.contact.type"].sudo().search([])
    providers = (
        request.env["res.partner"]
        .sudo()
        .search([("contact_type_id.type_of_contact", "=", "Supplier")])
    )
    customers = (
        request.env["res.partner"]
        .sudo()
        .search([("contact_type_id.type_of_contact", "=", "Customer")])
    )
    register_type = kw.get("type", "accreditation")
    banner = operation_banner_img.get(
        register_type, operation_banner_img["accreditation"]
    )
    alimentos_records = (
        request.env["product.template"]
        .sudo()
        .search([("product_type", "=", "alimento"), ("de_importacion", "=", True)])
    )
    alimentos_de_importacion = json.dumps([{'id': p.id, 'name': p.name} for p in alimentos_records])

    payment_methods = request.env["pos.payment.method"].sudo().search([])
    tipos_envase = request.env["pyxel.tipo.envase"].sudo().search([])
    tipos_envase_json = json.dumps([{'id': te.id, 'name': te.name} for te in tipos_envase])

    # URL param ?selected=17,16 (from nomenclador redirect) takes priority over stale session
    selected_from_url = kw.get('selected', '')
    if selected_from_url:
        try:
            product_selected = [int(x) for x in selected_from_url.split(',') if x.strip().isdigit()]
        except Exception:
            product_selected = request.session.get("product_selected", [])
        from_nomenclador = True
    else:
        product_selected = request.session.get("product_selected", [])
        referer = request.httprequest.referrer or ''
        from_nomenclador = '/nomenclador' in referer

    initially_selected_products = alimentos_records.filtered(lambda p: p.id in product_selected)

    # Recuperar estado guardado del header (si el usuario vino del nomenclador)
    saved_header = {}
    saved_product_rows = {}  # dict: product_id (str) -> {cantidad, tipo_envase}
    if from_nomenclador:
        saved_header = request.session.pop('import_header_state', {})
        for row in saved_header.pop('product_rows', []):
            pid = str(row.get('product_id', ''))
            if pid:
                saved_product_rows[pid] = {
                    'cantidad': row.get('cantidad', ''),
                    'tipo_envase': row.get('tipo_envase', ''),
                }

    render_values = {
        "countries": country.get_website_sale_countries(),
        "states": request.env["res.country.state"].sudo().search([]),
        "contact_types": contact_types,
        "partner_id": request.env.user.partner_id.id,
        "providers": providers,
        "customers": customers,
        "banner": banner,
        "register_type": register_type,
        "registered_user": False,
        "productos_seleccionados_ids": product_selected,
        "alimentos_de_importacion": alimentos_de_importacion,
        "alimentos_list": alimentos_records,
        "initially_selected_products": initially_selected_products,
        "payment_methods": payment_methods,
        "from_nomenclador": from_nomenclador,
        "saved_header": saved_header,
        "saved_product_rows": saved_product_rows,
        "tipos_envase": tipos_envase,
        "tipos_envase_json": tipos_envase_json,
    }
    crm_lead_exists = (
        request.env["crm.lead"]
        .sudo()
        .search(
            [("partner_id", "=", request.env.user.commercial_partner_id.id)], limit=1
        )
    )
    render_values["crm_lead_exists"] = bool(crm_lead_exists)

    if "uid" in request.context:
        render_values["registered_user"] = True

    return render_values


def loged_in():
    user = request.env.user
    if user._is_public():
        return request.redirect("/web/login")


class WebsiteForm(form.WebsiteForm):

    @http.route("/check_duplicate_nit", type="json", auth="public", website=True)
    def check_duplicate_nit(self, nit, **kw):
        """Verifica si hay duplicados de NIT entre los clientes."""
        if not nit:
            return True
        partner = (
            request.env["res.partner"]
            .sudo()
            .search([("vat", "=", nit), ("is_company", "=", True)])
        )
        return not bool(partner)

    @http.route("/check_file_type", type="json", auth="public", website=True)
    def check_file_type(self, config_param, file_type, **kw):
        attachment_id_str = (
            request.env["ir.config_parameter"]
            .sudo()
            .get_param(f"{config_param}.attachment_id")
        )
        if attachment_id_str:
            try:
                attachment_id = int(attachment_id_str)
            except ValueError:
                return
            attachment = request.env["ir.attachment"].sudo().browse(attachment_id)
            if attachment and attachment.mimetype:
                return attachment.mimetype == file_type

    @http.route("/get_form_management_types", type="json", auth="public", website=True)
    def get_form_management_types(self, contact_type_id=None, **kw):
        res = {}
        if contact_type_id:
            contact_types = (
                request.env["res.partner.contact.type"]
                .sudo()
                .search([("id", "=", int(contact_type_id))])
            )
            res = {str(x.id): x.name for x in contact_types.management_type_ids}
        return res

    @http.route("/get_form_states", type="json", auth="public", website=True)
    def get_form_states(self, country_code=None, **kw):
        res = {}
        if country_code:
            states = (
                request.env["res.country.state"]
                .sudo()
                .search([("country_id.code", "=", country_code)])
            )
            res = {str(x.id): x.name for x in states}
        return res

    @http.route("/get_form_cities", type="json", auth="public", website=True)
    def get_form_cities(self, state_id=None, **kw):
        res = {}
        if state_id:
            states = (
                request.env["res.city"]
                .sudo()
                .search([("state_id", "=", int(state_id))])
            )
            res = {str(x.id): x.name for x in states}
        return res

    def _get_partner_data(self, kwargs):
        partner_data = {
            "name": kwargs.get("parent_company_name"),
            "vat": kwargs.get("nit", False),
            "dap": kwargs.get("dap", False),
            "company_type": "company",
            "phone": kwargs.get("phone"),
            "email": kwargs.get("parent_company_email"),
            "country_id": int(kwargs.get("country", request.env.ref("base.cu").id)),
            "state_id": int(kwargs.get("state", False)),
            "street": kwargs.get("address"),
            "license_holder": kwargs.get("license_holder"),
            "management_type_id": int(kwargs.get("fgne_type", False)),
            "deed_number": int(kwargs.get("deed_input_number", False)),
            "deed_date": kwargs.get("deed_input_date"),
            "contact_type_id": int(kwargs.get("contact_type", False)),
        }
        if kwargs.get("supplier_type"):
            if kwargs.get("supplier_type") == "Productor":
                category_id = request.env.ref(
                    "pyxel_import_backend.res_partner_category_producer"
                ).id
                partner_data["category_id"] = [(4, category_id)]
            elif kwargs.get("supplier_type") == "Comerciante":
                category_id = request.env.ref(
                    "pyxel_import_backend.res_partner_category_businessman"
                ).id
                partner_data["category_id"] = [(4, category_id)]

        if kwargs.get("city"):
            city_id = int(kwargs.get("city", False))
            partner_data["city"] = (
                request.env["res.city"]
                .sudo()
                .search([("id", "=", city_id)], limit=1)
                .name
            )
            partner_data["city_id"] = city_id
        return partner_data

    def website_form(self, model_name, **kwargs):
        tipo_registro = kwargs.get("register_type")
        public_user = request.env.user.sudo()

        if model_name == "crm.lead" and tipo_registro == "accreditation":
            crm_lead_exists = (
                request.env["crm.lead"]
                .sudo()
                .search(
                    [("partner_id", "=", request.env.user.commercial_partner_id.id)],
                    limit=1,
                )
            )
            if crm_lead_exists:
                return Response(
                    json.dumps(
                        {
                            "error": "Usted ya se ha acreditado, para volver a acreditarse debe hacerlo con un usuario nuevo que no esté acreditado"
                        }
                    ),
                    status=400,
                    headers={"Content-Type": "application/json"},
                )
            contact_type_id = kwargs.get("contact_type")
            is_foreign_client = False
            if contact_type_id:
                try:
                    # Fetching the contact type record using the ID submitted by the form
                    contact_type_rec = (
                        request.env["res.partner.contact.type"]
                        .sudo()
                        .browse(int(contact_type_id))
                    )
                    if (
                        contact_type_rec.exists()
                        and contact_type_rec.nationality_type == "foreign"
                    ):
                        is_foreign_client = True
                except (ValueError, TypeError):
                    pass
            # Only validate NIT if it's a national client AND a NIT was actually provided
            nit_value = kwargs.get("nit")
            if not is_foreign_client and nit_value:
                valid_nit = self.check_duplicate_nit(nit_value)
                if not valid_nit:
                    return Response(
                        json.dumps(
                            {
                                "error": "El NIT ingresado ya existe. Verifique la información antes de continuar"
                            }
                        ),
                        status=400,
                        headers={"Content-Type": "application/json"},
                    )

            partner_data = self._get_partner_data(kwargs)
            partner = request.env["res.partner"].sudo().create(partner_data)
            public_user.partner_id.write(
                {"name": kwargs["partner_name"], "parent_id": partner.id}
            )
            request.params.update(
                {
                    "partner_id": partner.id,
                    "email_from": kwargs.get("parent_company_email"),
                    "partner_name": kwargs.get("parent_company_name"),
                }
            )
            kwargs.update(
                {
                    "partner_id": partner.id,
                    "email_from": kwargs.get("parent_company_email"),
                    "partner_name": kwargs.get("parent_company_name"),
                }
            )

            # Archivos de documentación individual (doc_*) y adjuntos fijos
            file_keys = [key for key in kwargs.keys() if key.startswith("doc_")]

            if kwargs.get("ficha_cliente[0][0]"):
                file_keys.append("ficha_cliente[0][0]")
            if kwargs.get("planilla_proveedor[0][0]"):
                file_keys.append("planilla_proveedor[0][0]")
            if kwargs.get("cuban_partner[0][0]"):
                file_keys.append("cuban_partner[0][0]")

            for file_key in file_keys:
                file = kwargs.get(file_key)
                if not file or not hasattr(file, 'seek'):
                    continue
                file.seek(0)
                base_key = file_key.split('[')[0]
                attachment_name = DOC_FIELD_LABELS.get(base_key, file.filename)
                request.env["ir.attachment"].sudo().create(
                    {
                        "name": attachment_name,
                        "datas": base64.b64encode(file.read()),
                        "res_model": "res.partner",
                        "res_id": partner.id,
                        "type": "binary",
                        "mimetype": file.mimetype,
                    }
                )

        elif model_name == "sale.order" and tipo_registro == "import":
            # Parse per-row data (product + cantidad + tipo_envase)
            product_rows = []
            try:
                product_rows = json.loads(kwargs.pop('product_rows_json', '[]') or '[]')
            except Exception:
                pass
            # Fallback to legacy productRequired multi-select
            if not product_rows:
                legacy_ids = kwargs.get('productRequired', '')
                if legacy_ids:
                    if isinstance(legacy_ids, str):
                        legacy_ids = [legacy_ids]
                    product_rows = [{'product_id': pid, 'cantidad': 1, 'tipo_envase': ''} for pid in legacy_ids]

            # Store for use in insert_record
            request.session['product_rows_data'] = product_rows

            pid = public_user.partner_id
            partner_id_val = pid.parent_id.id if pid.parent_id else pid.id
            request.params.update({"partner_id": partner_id_val})
            kwargs.update({"partner_id": partner_id_val})
            request.params.pop("productRequired", None)
            request.params.pop("register_type", None)
            kwargs.pop("productRequired", None)
            kwargs.pop("register_type", None)

            if kwargs.get("commitment_date"):
                kwargs["commitment_date"] = kwargs["commitment_date"] + " 12:00:00"
                request.params["commitment_date"] = kwargs["commitment_date"]

            # Guardar campos del encabezado en sesión para insert_record
            # (extract_data de Odoo solo procesa campos whitelisted en el form builder)
            request.session['import_header_fields'] = {
                'forma_pago': kwargs.get('forma_pago'),
                'presupuesto_disponible': kwargs.get('presupuesto_disponible'),
                'observaciones_solicitud': kwargs.get('observaciones_solicitud'),
                'note': kwargs.get('note'),
                'commitment_date': kwargs.get('commitment_date'),
            }

        res = super(WebsiteForm, self).website_form(model_name, **kwargs)
        _logger.info("Todas las claves disponibles en kwargs: %s", kwargs.keys())

        request.session["product_selected"] = []

        return res

    def insert_record(self, request, model, values, custom, meta=None):
        is_lead_model = model.model == "crm.lead"
        if is_lead_model:
            visitor_sudo = request.env["website.visitor"]._get_visitor_from_request()

            if "company_id" not in values:
                values["company_id"] = request.website.company_id.id
            lang = request.context.get("lang", False)
            values["lang_id"] = values.get("lang_id") or request.env[
                "res.lang"
            ]._lang_get_id(lang)

        if model.model == "sale.order":
            values["user_id"] = None
            product_rows_data = request.session.get('product_rows_data', [])
            order_lines = []
            for row in product_rows_data:
                try:
                    tmpl_id_int = int(row.get('product_id', 0))
                    qty = float(row.get('cantidad', 1) or 1)
                except Exception:
                    continue
                if not tmpl_id_int:
                    continue
                prod = request.env["product.product"].sudo().search(
                    [("product_tmpl_id", "=", tmpl_id_int)], limit=1
                )
                if prod:
                    order_lines.append((0, 0, {
                        "product_id": prod.id,
                        "product_uom_qty": qty,
                        "tipo_envase": int(row.get('tipo_envase', 0) or 0) or False,
                    }))
            values["order_line"] = order_lines
            values.pop("productRequired", None)
            values.pop("product_rows_json", None)

            # Aplicar campos del encabezado desde sesión (bypassea whitelist de extract_data)
            header = request.session.pop('import_header_fields', {})
            if header.get('forma_pago'):
                try:
                    values['forma_pago'] = int(header['forma_pago'])
                except (ValueError, TypeError):
                    pass
            if header.get('presupuesto_disponible'):
                try:
                    values['presupuesto_disponible'] = float(header['presupuesto_disponible'])
                except (ValueError, TypeError):
                    pass
            for fld in ('observaciones_solicitud', 'note'):
                if header.get(fld):
                    values[fld] = header[fld]
            cd = header.get('commitment_date') or ''
            if cd:
                values['commitment_date'] = cd

        if is_lead_model or model.model == "sale.order":
            record = (
                request.env[model.model]
                .sudo()
                .with_context(
                    mail_create_nosubscribe=True,
                )
                .create(values)
            )
            if custom or meta:
                _custom_label = "%s\n___________\n\n" % _(
                    "Other Information:"
                )  # Title for custom fields
                default_field = model.website_form_default_field_id
                default_field_data = values.get(default_field.name, "")
                custom_content = (
                    (default_field_data + "\n\n" if default_field_data else "")
                    + (_custom_label + custom + "\n\n" if custom else "")
                    + (self._meta_label + "\n________\n\n" + meta if meta else "")
                )

                # If there is a default field configured for this model, use it.
                # If there isn't, put the custom data in a message instead
                if default_field.name:
                    if default_field.ttype == "html":
                        custom_content = nl2br_enclose(custom_content)
                    record.update({default_field.name: custom_content})

            result = record.id

            if model.model == "sale.order":
                try:
                    import io
                    import datetime as dt
                    from docx import Document

                    attachment_id_str = (
                        request.env["ir.config_parameter"]
                        .sudo()
                        .get_param("solicitud.attachment_id")
                    )
                    if attachment_id_str:
                        tmpl = request.env["ir.attachment"].sudo().browse(int(attachment_id_str))
                        if tmpl.exists() and tmpl.datas:
                            doc = Document(io.BytesIO(base64.b64decode(tmpl.datas)))
                            table = doc.tables[0]

                            # Fila 0: una línea por producto con cantidad y tipo de envase
                            product_lines = []
                            for ol in record.order_line:
                                if not ol.product_id:
                                    continue
                                qty = int(ol.product_uom_qty) if ol.product_uom_qty == int(ol.product_uom_qty) else ol.product_uom_qty
                                tipo = ol.tipo_envase.name if ol.tipo_envase else ''
                                parts = [ol.product_id.name, str(qty)]
                                if tipo:
                                    parts.append(tipo)
                                product_lines.append(' — '.join(parts))
                            productos = '\n'.join(product_lines)

                            if record.commitment_date:
                                from pytz import timezone, utc as pytz_utc
                                user_tz_name = request.env.user.tz or 'UTC'
                                try:
                                    local_dt = record.commitment_date.replace(tzinfo=pytz_utc).astimezone(timezone(user_tz_name))
                                    fecha_entrega = local_dt.strftime('%d/%m/%Y')
                                except Exception:
                                    fecha_entrega = record.commitment_date.strftime('%d/%m/%Y')
                            else:
                                fecha_entrega = ''
                            presupuesto = (
                                f"{record.presupuesto_disponible:.3f} USD"
                                if record.presupuesto_disponible else ''
                            )
                            row_values = [
                                productos,       # Fila 0: producto — cantidad — envase
                                '',              # Fila 1: cantidad (va en fila 0)
                                fecha_entrega,   # Fila 2
                                '',              # Fila 3: tipo envase (va en fila 0)
                                record.forma_pago.name if record.forma_pago else '',
                                presupuesto,
                            ]

                            for i, val in enumerate(row_values):
                                table.rows[i].cells[1].paragraphs[0].add_run(val)

                            # Especificaciones (note es campo HTML, hay que limpiar etiquetas)
                            note_raw = record.note or ''
                            if '<' in note_raw:
                                note_raw = re.sub(r'<br\s*/?>', '\n', note_raw)
                                note_raw = re.sub(r'</p>', '\n', note_raw)
                                note_raw = re.sub(r'<[^>]+>', '', note_raw)
                                note_raw = html_lib.unescape(note_raw).strip()
                            table.rows[6].cells[0].add_paragraph(note_raw)
                            table.rows[7].cells[0].add_paragraph(record.observaciones_solicitud or '')
                            # Nombre del cliente: usar el partner del usuario logueado
                            user_partner = request.env.user.partner_id
                            empresa_cli = user_partner.parent_id.name if user_partner.parent_id else ''
                            usuario_cli = user_partner.name or ''
                            if empresa_cli:
                                table.rows[8].cells[0].add_paragraph(empresa_cli)
                                table.rows[8].cells[0].add_paragraph(usuario_cli)
                            else:
                                table.rows[8].cells[0].add_paragraph(usuario_cli)

                            today = dt.date.today()
                            table.rows[11].cells[3].paragraphs[0].add_run(str(today.day).zfill(2))
                            table.rows[11].cells[4].paragraphs[0].add_run(str(today.month).zfill(2))
                            table.rows[11].cells[5].paragraphs[0].add_run(str(today.year))

                            out = io.BytesIO()
                            doc.save(out)
                            out.seek(0)
                            request.env["ir.attachment"].sudo().create({
                                'name': f'Solicitud_Importacion_{record.partner_id.name}.docx',
                                'type': 'binary',
                                'datas': base64.b64encode(out.getvalue()).decode('utf-8'),
                                'res_model': 'sale.order',
                                'res_id': record.id,
                                'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            })
                except Exception as e:
                    _logger.warning("No se pudo generar el DOCX de solicitud: %s", e)

        else:
            # Llama al método insert_record del website, no del website_crm
            result = super(WebsiteForm2, self).insert_record(
                request, model, values, custom, meta=meta
            )

        if is_lead_model and visitor_sudo and result:
            lead_sudo = request.env["crm.lead"].browse(result).sudo()
            if lead_sudo.exists():
                vals = {"lead_ids": [(4, result)]}
                if not visitor_sudo.lead_ids and not visitor_sudo.partner_id:
                    vals["name"] = lead_sudo.contact_name
                visitor_sudo.write(vals)
        return result


class ControllerTest(http.Controller):

    @http.route("/business-register/save_header_state", type="json", auth="user", website=True, methods=["POST"])
    def save_header_state(self, **kwargs):
        product_rows = []
        try:
            product_rows = json.loads(kwargs.get('_product_rows', '[]') or '[]')
        except Exception:
            pass
        request.session['import_header_state'] = {
            'commitment_date': kwargs.get('commitment_date', ''),
            'forma_pago': kwargs.get('forma_pago', ''),
            'presupuesto_disponible': kwargs.get('presupuesto_disponible', ''),
            'note': kwargs.get('note', ''),
            'observaciones_solicitud': kwargs.get('observaciones_solicitud', ''),
            'product_rows': product_rows,
        }
        return {'ok': True}

    @http.route("/business-register/update_session_products", type="json", auth="user")
    def actualizar_sesion(self, selected_products, action=None):
        """
        Actualiza la variable de sesión `product_selected` con los valores seleccionados.
        Puede agregar o eliminar productos según el parámetro `action`. Si no se le pasa por param, por defecto agrega.
        """
        action = action or "add"

        if not isinstance(selected_products, list):
            selected_products = []
        else:
            selected_products = [
                int(p)
                for p in selected_products
                if isinstance(p, (int, str)) and str(p).isdigit()
            ]

        # Obtén la lista de productos previamente seleccionados de la sesión
        previous_products = request.session.get("product_selected", [])

        # Asegúrate de que previous_products sea una lista válida
        if not isinstance(previous_products, list):
            previous_products = []
        else:
            # Convierte todos los elementos a enteros si son válidos
            previous_products = [
                int(p)
                for p in previous_products
                if isinstance(p, (int, str)) and str(p).isdigit()
            ]
        if action == "add":
            updated_products = list(set(previous_products + selected_products))
        elif action == "remove":
            updated_products = [p for p in previous_products if p not in selected_products]
        elif action == "replace":
            updated_products = selected_products
        else:
            return {"status": "error", "message": f"Acción no válida: {action}"}

        # Actualizar la sesión con los nuevos productos seleccionados
        request.session["product_selected"] = updated_products
        request.session.modified = True

        return {"status": "success", "message": "Sesión actualizada correctamente"}

    @http.route("/business-register-thanks", type="http", auth="public", website=True)
    def business_register_thanks(self, **kw):
        crm_lead_exists = (
            request.env["crm.lead"]
            .sudo()
            .search(
                [("partner_id", "=", request.env.user.commercial_partner_id.id)],
                limit=1,
            )
        )
        days_in_process = "False"

        if crm_lead_exists:
            is_accredited = crm_lead_exists.partner_id.is_accredited

            if not is_accredited:
                days_in_process = (datetime.now() - crm_lead_exists.create_date).days

        return request.render(
            "pyxel_import_website.business_register_thanks",
            {"days_in_process": days_in_process},
        )

    @http.route("/business-register", type="http", auth="public", website=True)
    def controller_register(self, **kw):
        if request.env.user.id == request.env.ref("base.public_user").id:
            return request.redirect(
                f"/web/login?redirect=/business-register?type={kw.get('type', 'accreditation')}"
            )

        crm_lead_exists = (
            request.env["crm.lead"]
            .sudo()
            .search(
                [("partner_id", "=", request.env.user.commercial_partner_id.id)],
                limit=1,
            )
        )
        is_accredited = False

        if crm_lead_exists:
            is_accredited = crm_lead_exists.partner_id.is_accredited
        if kw.get("type") == "accreditation" and crm_lead_exists:
            if is_accredited:
                return request.redirect("/my/home")
            else:
                return request.redirect("/business-register-thanks")

        # Si llega al formulario de importación sin venir del nomenclador (sin ?selected=),
        # limpiar sesión para evitar que se muestren datos de pruebas anteriores
        if kw.get("type") == "import" and not kw.get("selected"):
            referer = request.httprequest.referrer or ''
            if '/nomenclador' not in referer:
                request.session['product_selected'] = []
                request.session.pop('import_header_fields', None)

        render_values = get_render_values(kw)

        if kw.get("type") == "import":
            # Si no se ha realizado el formulario de acreditación
            if not crm_lead_exists:
                return request.render(
                    "pyxel_import_website.waiting_for_active_contract"
                )
            # Si no es Cliente nacional no puede solicitar una importación
            if (
                request.env.user.commercial_partner_id.contact_type_id.type_of_contact
                == "Client"
                and request.env.user.commercial_partner_id.contact_type_id.nationality_type
                == "national"
            ):
                pass
            else:
                return request.render(
                    "pyxel_import_website.you_are_not_a_national_client",
                    {
                        "contact_type": request.env.user.commercial_partner_id.contact_type_id.name
                    },
                )
            # Si realizó el formulario de acreditación pero no está acreditado
            if not is_accredited:
                return request.redirect("/business-register-thanks")
            return request.render(
                "pyxel_import_website.import_registration", render_values
            )

        return request.render(
            "pyxel_import_website.business_registration", render_values
        )


class PreviewSolicitudController(http.Controller):

    @http.route("/business-register/preview-solicitud", type="http", auth="user", methods=["POST"], csrf=False)
    def preview_solicitud(self, **kwargs):
        import io
        import datetime as dt
        from docx import Document as DocxDocument

        attachment_id_str = (
            request.env["ir.config_parameter"].sudo().get_param("solicitud.attachment_id")
        )
        if not attachment_id_str:
            return request.redirect("/business-register?type=import")

        tmpl = request.env["ir.attachment"].sudo().browse(int(attachment_id_str))
        if not tmpl.exists() or not tmpl.datas:
            return request.redirect("/business-register?type=import")

        tipo_map = {
            'isotanque_20': 'Isotanque 20 pies',
            'isotanque_40': 'Isotanque 40 pies',
            'ibc': 'IBC',
        }

        forma_pago_name = ''
        forma_pago_id = kwargs.get('forma_pago')
        if forma_pago_id:
            try:
                pm = request.env['pos.payment.method'].sudo().browse(int(forma_pago_id))
                if pm.exists():
                    forma_pago_name = pm.name
            except Exception:
                pass

        fecha_entrega = ''
        commitment_date = kwargs.get('commitment_date', '')
        if commitment_date:
            try:
                fecha_entrega = datetime.strptime(commitment_date[:10], '%Y-%m-%d').strftime('%d/%m/%Y')
            except Exception:
                fecha_entrega = commitment_date

        presupuesto = ''
        presupuesto_val = kwargs.get('presupuesto_disponible')
        if presupuesto_val:
            try:
                presupuesto = f"{float(presupuesto_val):.3f} USD"
            except Exception:
                presupuesto = presupuesto_val

        # Parse product rows from form JSON; fallback to session
        product_rows = []
        try:
            product_rows = json.loads(kwargs.get('product_rows_json', '[]') or '[]')
        except Exception:
            pass
        if not product_rows:
            session_ids = request.session.get('product_selected', [])
            if session_ids:
                product_rows = [{'product_id': pid, 'cantidad': '', 'tipo_envase': ''} for pid in session_ids]

        # Cargar tipos de envase para lookup por ID
        _tipos_map = {
            str(te.id): te.name
            for te in request.env['pyxel.tipo.envase'].sudo().search([])
        }
        productos_lines = []
        for row in product_rows:
            pid = row.get('product_id', 0)
            if not pid:
                continue
            try:
                prod = request.env['product.product'].sudo().browse(int(pid))
                if not prod.exists():
                    continue
                line = prod.name
            except Exception:
                continue
            cant = str(row.get('cantidad', '')).strip()
            tipo = _tipos_map.get(str(row.get('tipo_envase', '')), '')
            if cant:
                line += f' — {cant}'
            if tipo:
                line += f' — {tipo}'
            productos_lines.append(line)
        productos = '\n'.join(productos_lines)

        partner = request.env.user.partner_id
        empresa_nombre = partner.parent_id.name or ''
        usuario_nombre = partner.name or ''

        doc = DocxDocument(io.BytesIO(base64.b64decode(tmpl.datas)))
        table = doc.tables[0]

        row_values = [
            productos,
            '',
            fecha_entrega,
            '',
            forma_pago_name,
            presupuesto,
        ]
        for i, val in enumerate(row_values):
            table.rows[i].cells[1].paragraphs[0].add_run(val)

        # Especificaciones (note puede venir como HTML del editor)
        note_raw = kwargs.get('note', '') or ''
        if '<' in note_raw:
            note_raw = re.sub(r'<br\s*/?>', '\n', note_raw)
            note_raw = re.sub(r'</p>', '\n', note_raw)
            note_raw = re.sub(r'<[^>]+>', '', note_raw)
            note_raw = html_lib.unescape(note_raw).strip()
        table.rows[6].cells[0].add_paragraph(note_raw)
        table.rows[7].cells[0].add_paragraph(kwargs.get('observaciones_solicitud', '') or '')
        if empresa_nombre:
            table.rows[8].cells[0].add_paragraph(empresa_nombre)
            table.rows[8].cells[0].add_paragraph(usuario_nombre)
        else:
            table.rows[8].cells[0].add_paragraph(usuario_nombre)

        today = dt.date.today()
        table.rows[11].cells[3].paragraphs[0].add_run(str(today.day).zfill(2))
        table.rows[11].cells[4].paragraphs[0].add_run(str(today.month).zfill(2))
        table.rows[11].cells[5].paragraphs[0].add_run(str(today.year))

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)

        filename = f'Solicitud_Importacion_{usuario_nombre}.docx'
        headers = [
            ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            ('Content-Disposition', f'attachment; filename="{filename}"'),
        ]
        return request.make_response(out.getvalue(), headers=headers)


class ProductSearchController(http.Controller):

    @http.route(
        ["/nomenclador", "/nomenclador/page/<int:page>"],
        type="http",
        auth="public",
        website=True,
    )
    def nomenclador_view(self, search=None, page=1, **kwargs):
        loged_in()
        """Renderiza la vista con el buscador y los resultados paginados."""
        alimentos_de_importacion = (
            request.env["product.template"]
            .sudo()
            .search([("product_type", "=", "alimento"), ("de_importacion", "=", True)])
        )
        filters = [("product_type", "=", "alimento"), ("de_importacion", "=", True)]
        domain = [("name", "ilike", search)] if search else []
        domain += filters

        from_view = kwargs.get("from", None)

        total_products = request.env["product.template"].sudo().search_count(domain)

        base_url = "/nomenclador"

        url_args = {}
        if from_view:
            url_args["from"] = from_view
        if search:
            url_args["search"] = search

        pager = request.website.pager(
            url=base_url,
            url_args=url_args,
            total=total_products,
            page=page,
            step=10,  # 10 productos por página
            scope=3,
        )

        products = (
            request.env["product.template"]
            .sudo()
            .search(domain, limit=10, offset=(page - 1) * 10)
        )

        return request.render(
            "pyxel_import_website.nomenclador_template",
            {
                "products": products,
                "search": search,
                "from_view": from_view,
                "alimentos_de_importacion": alimentos_de_importacion,
                "pager": pager,
                "selected_product_ids": request.session.get("product_selected", []),
            },
        )


class DownloadFileController(http.Controller):

    @http.route("/descargar/<string:file_name>", type="http", auth="public")
    def download_file(self, file_name, **kw):
        if file_name not in [
            "load_products",
            "solicitud",
            "ficha_cliente_estatal",
            "ficha_cliente_fgne_tcp",
            "perfil_proveedor",
            "cuban_partner",
        ]:
            return request.redirect("/web")
        attachment_id_str = (
            request.env["ir.config_parameter"]
            .sudo()
            .get_param(f"{file_name}.attachment_id")
        )
        if attachment_id_str:
            try:
                attachment_id = int(attachment_id_str)
            except ValueError:
                return request.redirect("/web")
            attachment = request.env["ir.attachment"].sudo().browse(attachment_id)
            if attachment and attachment.datas:
                return Stream.from_attachment(attachment).get_response(
                    as_attachment=True
                )
        return request.redirect("/web")
